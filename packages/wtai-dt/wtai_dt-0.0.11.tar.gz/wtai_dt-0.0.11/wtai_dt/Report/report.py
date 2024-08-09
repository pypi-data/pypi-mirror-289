#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
class Report:
    def __init__(self, cover_path=None):
        self.cover_path = cover_path

    def create_new_document(self):
        if self.cover_path:
            return Document(self.cover_path)
        return Document()

    def insert_lettle_title_word(self, doc, lettle_text):
        subtitle = doc.add_heading(level=2)
        subtitle_run = subtitle.add_run(lettle_text)
        subtitle_run.font.bold = True
        subtitle_run.font.name = '宋体'
        subtitle_run.font.color.rgb = RGBColor(0, 0, 0)
        subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        subtitle_run.font.size = Pt(12)

    def insert_paragraph(self, doc, text, font_name=u'宋体', font_size=12):
        paragraph = doc.add_paragraph('\t' + text)
        for run in paragraph.runs:
            self.set_font(run, font_name, font_size)

    def insert_image_word(self, doc, image_path):
        section = doc.sections[-1]
        page_width = section.page_width
        left_margin = section.left_margin
        right_margin = section.right_margin
        available_width = page_width - left_margin - right_margin
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(image_path, width=available_width)

    def set_cell_properties(self, cell):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def insert_table_word(self, doc, data):
        table = doc.add_table(rows=1, cols=len(data[0]))
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        for i, header in enumerate(data[0]):
            hdr_cells[i].text = header
            self.set_cell_properties(hdr_cells[i])

        for row in data[1:]:
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = value
                self.set_cell_properties(row_cells[i])

    def set_font(self, run, font_name=u'宋体', font_size=12):
        font = run.font
        font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        font.size = Pt(font_size)

    def insert_centered_title(self,paragraph, title):
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run(title)
        self.set_font(run)

    def insert_title(self, doc, title_text):
        title = doc.add_heading(level=1)
        title_run = title.add_run(title_text)
        title_run.font.bold = True
        title_run.font.name = '宋体'
        title_run.font.color.rgb = RGBColor(0, 0, 0)  # 设置为黑色字体
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        title_run.font.size = Pt(12)

    def change_colums_value(self, row, change, columns_list):
        if row in columns_list:
            index = columns_list.index(row)
            columns_list[index] = change
        return columns_list

    def save(self, doc, path):
        doc.save(path)

    def set_cell_borders(self, cell, pt=0):
        """ Set the borders of the cell to be invisible or visible """
        # Define the border properties: Set the border to no border
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = OxmlElement('w:tcBorders')
        for border_name in ('top', 'left', 'right', 'bottom'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tc_borders.append(border)
        tc_pr.append(tc_borders)

    def insert_empty_paragraph(self, paragraph):
        paragraph.add_run()

    def insert_image_table(self, doc, image_paths, image_name, num_cols_per_table=2):
        num_images = len(image_paths)
        num_tables = (num_images + num_cols_per_table - 1) // num_cols_per_table

        for i in range(num_tables):
            num_rows = 2
            table = doc.add_table(rows=num_rows, cols=num_cols_per_table)
            table.autofit = True
            table.style = 'Table Grid'
            table.allow_autofit = True
            table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            for row in table.rows:
                for cell in row.cells:
                    self.set_cell_borders(cell)

            for j in range(num_cols_per_table):
                cell_title = table.cell(0, j)
                cell_image = table.cell(1, j)

                image_index = i * num_cols_per_table + j
                if image_index < num_images:
                    self.insert_centered_title(cell_title.paragraphs[0], image_name[image_index])
                    paragraph = cell_image.paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = paragraph.add_run()
                    cell_width = table.cell(1, j).width - Cm(0.8)
                    run.add_picture(image_paths[image_index], width=cell_width)
                else:
                    self.insert_centered_title(cell_title.paragraphs[0], '')
                    self.insert_empty_paragraph(cell_image.paragraphs[0])

