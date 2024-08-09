from .utils import transliterate
from .mail_sender import Mail_sender

from docx import Document
from docx.shared import Inches

from IPython.display import clear_output

import pandas as pd
import torch
import numpy as np

import matplotlib.pyplot as plt

import os
import re
import difflib # смотреть расстояние между текстами



class Obraz(Mail_sender): 
  def __init__(self, obraz=None, lang=None, book=None, adress_book=None, model_chekpoint='sergiyvl/model_65000_20ep'):
    """
    Класс для работы с образом.
    obraz - str, сам образ. Обязателен при создании объекта. Образ менять нельзя (не имеет смысла).
    book - [" ", " "]можно задать здесь, а можно добавить отдельно.
    adress_book - ПОТЕНЦИАЛЬНО возможность по адресу (относительному или ссылки) загружать книгу.

    Методы:
    analize - основная функция. в ней происходит анализ book.
    """
    super().__init__("smtp.mail.ru", "le_i_van@mail.ru", "9dx-U7V-KJr-b2H")
    if lang == None: 
      lang = input("Если вы хотите работать с текстами на церковно-славянском языке, введите \"cs\"\n"+
                   "Если вы хотите работать с текстами на древнегреческом языке, введите \"greek\"\n"+
                   "Введите здесь: ")
      lang = lang.strip().upper()
    if lang.strip().upper() == "greek".upper():
      book = "all_alt_greek_15082023.csv"
      self.lang = "greek"
    else:
      book = "mto_148000_with_adress_newtriodi_23082023.csv"
      self.lang = "cs"
    
    self._wholebook = None
    self.book = None
    self.email = None
    self.model_chekpoint = model_chekpoint
    self.model = None
    self.tokenizer = None
    self.key_tokens = None # ключевые токены. массив индексов
    self.key_symbols = None # ключевые буквосочетания. массив строк
    self._book_helpfull = None
    self.texts_with_key_key = None # лучшие результаты, где встречаются ключевые буквосочетания
    self.texts_without_key_key = None # лучшие результаты, где НЕТ ключевых буквосочетаний
    self._fast_print_book = None # массив со склеинными похожими текстами

    if adress_book != None:
      self.add_book(adress_book=adress_book)
    elif book != None:
      self.add_book(book=book)

#########################################################################################################################################################
################################################### функциии для поиска, которые видит пользователь (дальше есть блок, где реализована техническая сторона процессов)

  def analize(self):
    """
    Основная функция. Запуская ее Вы анализируете весь book, который добавили в последний раз. Если вы еще не добавили, то вас попросят это сделать.
    """
    if self.book == None:
      raise ValueError("Добавьте книгу для анализа. Для этого можно воспользоваться методом .add_book()")


    if self.model == None or self.tokenizer == None:
      self.set_model_and_tokenizer()

    print("Ваш образ: \""+self.obraz+"\"")
    if self.key_tokens == None:
      self.set_key_tokens()
    if self.key_symbols == None:
      self.set_key_symbols()

    print("Анализируем... Анализ может занимать до 15 минут.")
    self._result = self._analize()

    self._make_book_helpfull(self._result)

    self.texts_with_key_key = []
    self.texts_without_key_key = []
    for text in self._book_helpfull:
        key_key_exist = True
        for key_key in self.key_symbols:
            if key_key.upper() not in text[0].upper():
                key_key_exist = False
        if key_key_exist:
            self.texts_with_key_key.append(text)
        else:
            self.texts_without_key_key.append(text)

    i = 0
    print("Лучшие тексты с ключевыми символами. Эти тексты можно снова увидеть, если вызвать метод .results(): ")
    for text in self.texts_with_key_key:
        i += 1
        if i < 40:
          print(i, text)

    # self.send_to_email()

    self.working_with_bookhelpfull()

################################################### всевозможный быстрый поиск

  def fast_or(self):
    """ В self.book проводится поиск буквосочетаний и тексты, где есть буквосочетания, те тексты оставляются. Остальные - удаляются
        Эта функция без использования регулярных выражений и без возможности поиска в начале, середине или конце слова
    """
    keys = self._input_for_fast()
    i = 0
    self._fast_book = []
    for text in self.book:
      y = False
      for key in keys:
        if key.upper() in text[1].upper():
          y = True
        # for word in text.split():
        #   if key.upper() in word.upper():
        #     print(word)
      if y:
        i += 1
        self._fast_book.append(text)
    self._fast_print_book = self._glue_together_similar_texts(self._fast_book)
    self.book = self._fast_book
    print("В библиотеке ", i, " текстов, где есть хотя бы одно из буквосочетаний ", str(keys))
    self._fast_print(keys, 'or')

  def fast_and(self):
    """ В self.book проводится поиск буквосочетаний и тексты, где есть буквосочетания, те тексты оставляются. Остальные - удаляются
        Эта функция без использования регулярных выражений и без возможности поиска в начале, середине или конце слова
    """
    keys = self._input_for_fast()
    i = 0
    self._fast_book = []
    for text in self.book:
      y = True
      for key in keys:
        if key.upper() not in text[1].upper():
          y = False
        # for word in text.split():
        #   if key.upper() in word.upper():
        #     print(word)
      if y:
        i += 1
        self._fast_book.append(text)
    self._fast_print_book = self._glue_together_similar_texts(self._fast_book)
    self.book = self._fast_book
    print("В библиотеке ", i, " текстов, где в каждом есть все буквосочетания ", str(keys))
    self._fast_print(keys, 'and')

  def fast_or_re(self):
    """ В self.book проводится поиск буквосочетаний и тексты, где есть буквосочетания, те тексты оставляются. Остальные - удаляются
        Особенность этой функция в использовании регулярных выражений и в возможности поиска в начале, середине или конце слова
    """
    keys = self._input_for_fast()
    i = 0
    self._fast_book = []
    pattern = r''
    for key in keys:
      key = key.upper()
      if key[0] == "_":
        key = r'[^\w]'+key[1:]
      if key[-1] == "_":
        key = key[:-1]+r'[^\w]'
      key += '|'
      pattern += key
    pattern = pattern[:-1]

    for text in self.book:
      if re.search(pattern, " "+text[1].upper()+" "):
        i += 1
        self._fast_book.append(text)
    self._fast_print_book = self._glue_together_similar_texts(self._fast_book)
    self.book = self._fast_book
    print("В библиотеке ", i, " текстов, где есть хотя бы одно из буквосочетаний ", str(keys))
    self._fast_print(keys, 'or')


  def fast_and_re(self):
    """ В self.book проводится поиск буквосочетаний и тексты, где есть буквосочетания, те тексты оставляются. Остальные - удаляются
        Особенность этой функция в использовании регулярных выражений и в возможности поиска в начале, середине или конце слова
    """
    keys = self._input_for_fast()
    i = 0
    self._fast_book = []
    patterns = []
    for key in keys:
      key = key.upper()
      if key[0] == "_":
        key = r'[^\w]'+key[1:]
      if key[-1] == "_":
        key = key[:-1]+r'[^\w]'
      patterns.append(key)

    for text in self.book:
      y = True
      for pattern in patterns:
        if not re.search(pattern, " "+text[1].upper()+" "):
          y = False
      if y:
        i += 1
        self._fast_book.append(text)
    self._fast_print_book = self._glue_together_similar_texts(self._fast_book)
    self.book = self._fast_book
    print("В библиотеке ", i, " текстов, где в каждом есть все буквосочетания ", str(keys))
    self._fast_print(keys, 'and')


#########################################################################################################################################################
################################################### полезные функции, которые пока не удалось разбить на блоки

  def results(self):
    """
    Функция для вывода результатов поиска.
    """
    i = 0
    for text in self.texts_with_key_key[:30]:
      i += 1
      print(i, text)
    i = 0
    for text in self.texts_without_key_key[:30]:
      i += 1
      print(i, text)


  def add_book(self, book=None, adress_book=None):
    """Добавить книгу для поиска. Если не указан адрес - подгружает книгу по умолчанию.
    """

    if book != None:
      p = os.path.join(os.path.dirname(__file__), '')
      p = os.path.join(p, book)
      csv_book = pd.read_csv(p)

      self._wholebook = [sent[1:4] for sent in csv_book.values]
      for i in range(len(self._wholebook)):
        if type(self._wholebook[i][0]) != type(148.97):
          self._wholebook[i][0] = re.sub(r'<.*?>', '', self._wholebook[i][0])
          self._wholebook[i][0] = re.sub(r'[\n\t]', '', self._wholebook[i][0])
          self._wholebook[i][0] = self._wholebook[i][0].strip()
          self._wholebook[i][1] = re.sub(r'<.*?>', '', self._wholebook[i][1])
          self._wholebook[i][1] = re.sub(r'[\n\t]', '', self._wholebook[i][1])
          self._wholebook[i][1] = self._wholebook[i][1].strip()
          self._wholebook[i][2] = re.sub(r'<.*?>', '', self._wholebook[i][2])
          self._wholebook[i][2] = re.sub(r'[\n\t]', '', self._wholebook[i][2])
          self._wholebook[i][2] = self._wholebook[i][2].strip()
        else:
          self._wholebook[i][0] = "пустое место"
      self.book = self._wholebook

    elif adress_book != None:
      csv_book = pd.read_csv(adress_book)

      self._wholebook = [sent[1:4] for sent in csv_book.values]
      for i in range(len(self._wholebook)):
        if type(self._wholebook[i][0]) != type(148.97):
          self._wholebook[i][0] = re.sub(r'<.*?>', '', self._wholebook[i][0])
          self._wholebook[i][0] = re.sub(r'[\n\t]', '', self._wholebook[i][0])
          self._wholebook[i][0] = self._wholebook[i][0].strip()
          self._wholebook[i][1] = re.sub(r'<.*?>', '', self._wholebook[i][1])
          self._wholebook[i][1] = re.sub(r'[\n\t]', '', self._wholebook[i][1])
          self._wholebook[i][1] = self._wholebook[i][1].strip()
          self._wholebook[i][2] = re.sub(r'<.*?>', '', self._wholebook[i][2])
          self._wholebook[i][2] = re.sub(r'[\n\t]', '', self._wholebook[i][2])
          self._wholebook[i][2] = self._wholebook[i][2].strip()
        else:
          self._wholebook[i][0] = "пустое место"
      self.book = self._wholebook
      # print(self.book[:100])

  def send_to_email(self, email=None):
    """
    Реализована возможность отправить результаты на почту. На вход электронный адрес.
    """
    h_m_with_keykey = int(input("Сколько лучших результатов Вы хотите получить? Введите число: "))
    if self.email != None:
      email = self.email
    elif email == None:
      # inp = input("Вы хотите получить лучшие результаты на почту в виде word файла? Введите Да или Нет: ")
      # if "нет".upper() in inp.upper():
      #   return
      # if input("Хотите увидеть диаграмму распределения для текстов с ключевыми токенами? Введите 1 или 0 ") == "1":
      #   self.show_diagramm(self.texts_with_key_key)
      print("Исходя из опыта, советуем Вам указать gmail. yandex или outlook тоже хорошо. mail - плохо. Не забудьте проверить спам!")
      email = input("Введите адрес Вашей электронной почты: ")
    if self.texts_with_key_key == None:
      print("Чтобы результаты отправить, их нужно получить. Выполните анализ книги. .analize()")
      return
    inp = input("Вы хотите получить результаты в формате docx или просто внутри письма? Введите 'docx' или любые символы: ")
    if "docx" in inp:
      self._send_to_email_docx(email, h_m_with_keykey)
    else:
      self._send_to_email(email, h_m_with_keykey)


  def working_with_bookhelpfull(self):
    """
    Переход в режим работы с bookhelpfull.
    т е весь поиск будет теперь осуществляться в тех результатах полноценного поиска по всей книге, где есть хотя бы один образ.
    В будущем будет реализован более тонкий подбор границы.
    """
    print("\n\nРежим работы с bookhelpfull активирован. \n")
    self.book = self._book_helpfull


  def working_with_wholebook(self):
    """
    Переход в режим работы со всей книгой. т е весь поиск будет осуществляться по всей добавленной книге.
    """
    print("\n\nРежим работы со всей книгой. \n")
    self.book = self._wholebook

  def _glue_together_similar_texts(self, array_of_texts): # на вход нужен сам массив текстов вида: (текст, адрес текста)
    ans = input("Поиск произведен. Получилось "+ str(len(array_of_texts)) + " текстов. Мы можем среди получившихся текстов найти почти идентичные и идентичные и оставить только их адреса.\n" +
                "Учитывайте, что 10 текстов анализируются за пару секунд, 100 примерно за 40, а 500 за несколько минут. \n" +
                "Введите \"да\", если хотите так сделать, иначе введите любые символы: ")
    if ans.strip().upper() == "ДА": 
      table_of_simil = []
      for text_1 in array_of_texts:
        texts_simil = []
        for text_2 in array_of_texts:
          texts_simil.append(difflib.SequenceMatcher(None, text_1[1], text_2[1]).ratio()) #  real_quick_ratio() quick_ratio() ratio()
        table_of_simil.append(texts_simil)

      array_in_work = [{"text_el": array_of_texts[i], "base_id": i} for i in range(len(array_of_texts))]
      
      if len(table_of_simil) > 0:
        for i in range(len(table_of_simil[0])):
          for j in range(i, len(table_of_simil)):
            if i != j: # чтобы не сломалось!!
              if table_of_simil[i][j] >= 0.85:
                array_in_work[j]["base_id"] = array_in_work[i]["base_id"]

      simil_itog = [[] for _ in array_in_work]
      for el in array_in_work:
        if len(simil_itog[el["base_id"]]) == 0: # если длина массива под номером base_id в массиве simil_itog равна нулю, значит, в этот массив ничего не добавляли, значит, надо базово настроить
          simil_itog[el["base_id"]].append(el["text_el"][0])
          simil_itog[el["base_id"]].append(el["text_el"][1])
          simil_itog[el["base_id"]].append([])
        simil_itog[el["base_id"]][2].append(el["text_el"][2])

      simil_itog_itog = []
      for el in simil_itog:
        if len(el) > 0:
          a = [len(el[2]), el[0], el[1], el[2]]
          simil_itog_itog.append(a)
      return sorted(simil_itog_itog, reverse=True)
    else: 
      arr = []
      i = 0

      for el in array_of_texts: 
        i += 1
        arr.append([el[0], el[1], el[2]])
      return arr

  def _altgreek_diakr_to_lite(self, text):
    """ переводит строку греческих текстов с диакритикой в греческий текст без диакритики
    """
    dict_altgreek_diakr = {
      "???????????????????????????????????????????????": "?",
      "?????????????????": "?",
      "???????????????????????????????????????????": "?",
      "?????????????????????????????": "?",
      "?????????????????": "?",
      "???????????????????????????????????????????": "?",
      "?????????????????????????": "?",
      "?'???": "?",
      "?`???????????????": "",
      "?,.?": " "
    }

    text = list(text)
    for i in range(len(text)):
      for key in dict_altgreek_diakr.keys():
        if text[i] in key:
          text[i] = dict_altgreek_diakr[key]

    return "".join(text)

  def _input_for_fast(self):
    """выделен в отдельный модуль во избежание опечаток
    """
    text = input("Режим Fast. Введите через пробел те буквосочетания, которые должны быть в богослужебном тексте: ")
    if self.lang == "greek":
      text = self._altgreek_diakr_to_lite(text)
    return [key for key in text.split()]



#########################################################################################################################################################
################################################### set модуль. тут прописаны функции, которые задают ключвые моменты. например, сам образ и книгу для поиска

  def set_key_tokens(self):
    """
    Задать ключевые токены.
    """
    if self.tokenizer != None:
      obraz_tokenize = self.tokenizer.tokenize(self.obraz)

      i = 0
      for token in obraz_tokenize:
          i += 1
          print("(", i,", "+token+")", end="  ", sep="")
          if i%10 == 0:
              print()
      self.key_tokens = [int(i) for i in input("\nВведите через пробел номера ключевых токенов: ").split()]

  def set_key_symbols(self):
    """Задать ключевые буквосочетания."""
    self.key_symbols = [i for i in input("Введите буквосочетания, без которых в тексте точно нет образа.\nЕсли возможно выделить несколько, тогда введите через пробел:").split()]

  def set_model_and_tokenizer(self, model_chekpoint=None, model=None, tokenizer=None):
    print("Загружаем модель для анализа...")
    self.device = torch.device("cuda:0" if torch.cuda.is_available() else "cpu")

    if model_chekpoint != None:
      self.model_chekpoint = model_chekpoint

    if model == None:
      from transformers import AutoTokenizer
      self.tokenizer = AutoTokenizer.from_pretrained(self.model_chekpoint)

      from transformers import AutoModelForMaskedLM
      self.model = AutoModelForMaskedLM.from_pretrained(self.model_chekpoint)
    else:
      self.model = model
      self.tokenizer = tokenizer

    self.model.to(self.device)
    clear_output(True)
    clear_output(True)
    clear_output(True)

  def set_obraz(self, obraz):
    obraz = input("Введите образ: ")
    while True:
      try:
        if type(obraz) != type("str"):
          raise TypeError("В качестве образа необходимо передать строку (str).")
      except TypeError:
        print("В качестве образа необходимо передать строку (str).")
        obraz = input("Введите образ: ")
      else:
        break
    if obraz[-1] != ".":
      obraz += "."
    clear_output()
    print("Ваш образ: \""+obraz+"\"")

    return obraz

#########################################################################################################################################################
################################################### функции умного поиска с использованием нейронных сетей (пока что только русский вариант)

  def _analize(self):
    """
    Седьмая вариация. Нормализация по вектору образа. В сравнение идут все вектора кроме первого, точки и последнего.
    Ключевые токены возводятся в степень. Ключевые токены, образующие слова, зависят друг от друга. (otn = batch_max[i][ind-1]/batch_max[i][ind]).
    Есть возможность ввести ключевые буквосочетания.
    Добавлена грубая реакция на сочетание ключевых токенов. (коэфициент - сумма токенов > 200)
    """

    obraz_tokenize = self.tokenizer.tokenize(self.obraz)

    relevant = [(1 if i+1 not in self.key_tokens else 1.5) for i in range(len(obraz_tokenize))]


    # ТОКЕНИЗИРУЕМ ТЕКСТ И НОРМАЛИЗУЕМ СЛОВА
    batch_size=64
    text_tokenized = self.tokenizer(self.obraz, return_tensors='pt').to(self.device)
    text_throw_model = self.model(
                            **text_tokenized,
                            output_hidden_states=True
                        ).hidden_states[:][-1]


    words = text_throw_model[0][1:-1].reshape(1, -1, 768)

    normalization_token = []
    for word in obraz_tokenize:
        token = self.tokenizer(word+'.', return_tensors='pt').to(self.device)
        token_1 = self.model(
            **token,
            output_hidden_states=True
        ).hidden_states[:][-1]
        token_2 = torch.transpose(token_1, 1, 2)
        koef = 40/(token_1 @ token_2)[0, 1, 1].detach().cpu().numpy()
        normalization_token.append(koef)

    normalization_word = []
    # normilization = 40/(words@torch.transpose(words, 1, 2)) ЧТО ПО СУТИ ВНИЗУ ПРОИСХОДИТ
    for word in words[0]:
        word_transpose = torch.transpose(word.reshape(1, -1), 0, 1)
        koef = 40/(word.reshape(1, -1) @ word_transpose)[0, 0].detach().cpu().numpy()
        normalization_word.append(koef)

    if False:
      for i in range(len(normalization_word)):
          print(obraz_tokenize[i], normalization_token[i], normalization_word[i])

    # ИЩЕМ ПАРЫ КЛЮЧЕВЫХ ТОКЕНОВ, КОТОРЫЕ ОБРАЗУЮТ СЛОВА, ЧТОБЫ В БУДУЩЕМ НОРМАЛИЗОВАТЬ ОТНОСИТЕЛЬНО ИХ СОВМЕСТНОЙ ВСТРЕЧАЕМОСТИ
    pairs_second = []
    for i in range(len(relevant)):
        if relevant[i] > 1:
            if obraz_tokenize[i][:2] == "##":
                if relevant[i-1] > 1:
                    pairs_second.append(i)

    # ПРОГОНЯЕМ ЧЕРЕЗ МОДЕЛЬ И СОХРАНЯЕМ В result ИТОГИ
    result = np.zeros((len(self.book), 2, len(relevant)))

    self.local_book = [text[0] for text in self.book]

    for start_index in range(0, len(self.local_book), batch_size):
        batch = self.local_book[start_index:start_index+batch_size]

        batch = self.tokenizer(batch, return_tensors='pt',truncation=True, padding=True, max_length=45).to(self.device)

        batch = self.model(
        **batch,
        output_hidden_states=True
        ).hidden_states[:][-1]

    #  КАЖДЫЙ БАТЧ ПЕРЕМНОЖАЕТСЯ С МАТРИЦЕЙ ТЕКСТА И ЛУЧШИЕ РЕЗУЛЬТАТЫ ОСТАВЛЯЮТСЯ
        batch = torch.transpose(batch, 1, 2)

        batch = (words @ batch).detach().cpu().numpy()
        batch = batch[:, :, 1:-1]

        batch_max = np.max(batch, axis=2)
        for i in np.arange(batch_size):
            if start_index+i<len(self.local_book):
                result[start_index+i][0][0] = start_index+i
                for ind in pairs_second:
                  otn = batch_max[i][ind-1]/batch_max[i][ind]
                  normalization_word[ind-1] *= (otn if otn < 1 else 1/otn)

                result[start_index+i][1] = (batch_max[i]*normalization_word)**relevant # СМОТРИ КАКАЯ НОРМАЛИЗАЦИЯ
                # чем больше ключевых токенов тем лучше. ФУНКЦИЯ СТАЛА БОЛЕЕ СТУПЕНЧАТОЙ!!! ТЕПЕРЬ СТАЛО ПОНЯТНО, ДО КАКОЙ ДОБАВЛЯТЬ В bool_helpfull
                koef = sum((result[start_index+i][1] > 200)*1)
                result[start_index+i][1] *= (koef+1)/2

    return result


#########################################################################################################################################################
################################################### создание документов для последующей отправки по почте

  def make_result_dokument_not_docx(self, file_name, title=None, first_paragraf=None, h_m_with_keykey=70, h_m_without_keykey=20): #  adress = "/content/drive/MyDrive/diplom/results/results/",

    result_doc = """"""

    # document = Document()

    if title != None:
      result_doc += title.upper() + "\n" + "\n"

    # document.add_heading('Описание способа выбора', level=1)
    if first_paragraf != None:
      result_doc += first_paragraf + "\n"

    obraz_tokenized = self.tokenizer.tokenize(self.obraz)

    result_doc += "\n" + "\n" + "\n" + 'Лучшие тексты С ключевыми буквосочетаниями'.upper() + "\n" + "\n"
    result_doc += "Ключевые буквосочетания: " + str(self.key_symbols)
    i = 0
    for text in self.texts_with_key_key[:h_m_with_keykey]:
      i += 1
      result_doc += str(i)+" "+text + "\n"

    result_doc += "\n" + "\n" + "\n" + 'Лучшие тексты БЕЗ ключевых буквосочетаний'.upper()  + "\n" + "\n"
    i = 0
    for text in self.texts_without_key_key[:h_m_without_keykey]:
      i += 1
      result_doc += str(i)+" "+text + "\n"


    # document.save(adress + file_name)
    return result_doc


  def _make_result_dokument_docx(self, file_name, title=None, first_paragraf=None, h_m_with_keykey=70, h_m_without_keykey=20, adress = '/content/'): #  adress = '/content/',

    document = Document()

    if title != None:
      document.add_heading(title, 0)

    document.add_heading('Описание способа выбора', level=1)
    if first_paragraf != None:
      p = document.add_paragraph(first_paragraf)

    obraz_tokenized = self.tokenizer.tokenize(self.obraz)

    document.add_heading('Лучшие тексты С ключевыми буквосочетаниями', level=1)
    document.add_paragraph("Ключевые буквосочетания: " + str(self.key_symbols))

    i = 0
    num = input("Оставить нумерацию? введите Да или Нет").strip().upper()
    for text in self.texts_with_key_key[:h_m_with_keykey]:
      if num == "Нет".upper():
        p = document.add_paragraph(text[0]+"  ")
      else:
        i += 1
        p = document.add_paragraph(str(i)+" "+text[0]+"  ")
      p.add_run("("+text[2]+")").italic = True

    document.add_heading('Лучшие тексты БЕЗ ключевых буквосочетаний', level=1)
    i = 0
    for text in self.texts_without_key_key[:h_m_without_keykey]:
      i += 1
      p = document.add_paragraph(str(i)+" "+text[0]+"  ")
      p.add_run("("+text[2]+")").italic = True



    document.save(adress + file_name)


  def _make_fast_book_dokument_docx(self, file_name, keys, title=None, mode=None, adress = '/content/'): #  adress = '/content/',

    document = Document()

    if title != None:
      document.add_heading(title, 0)

    document.add_heading('Способ выбора текстов:', level=1)
    if mode == 'or':
      p = document.add_paragraph("Тексты, где есть хотя бы одно из буквосочетаний " + str(keys)[1:-1])
    elif mode == 'and':
      p = document.add_paragraph("Тексты, где в каждом есть все буквосочетания " + str(keys)[1:-1])

    i = 0
    num = input("Оставить нумерацию? введите Да или Нет").strip().upper()
    for text in self._fast_print_book:
      if num == "Нет".upper():
        p = document.add_paragraph(text[0]+"  ")
      else:
        i += 1
        p = document.add_paragraph(str(i)+". "+str(text[0])+" - "+text[1]+"   ")
      p.add_run("("+", ".join(text[2])+")").italic = True


    document.save(adress + file_name)




#########################################################################################################################################################
 ################################################### показ текстов и отправка по почте

  def _fast_print(self, keys, mode):
    choice = int(input("Хотите увидеть найденные тексты? Введите '0', если не хотите. \n" +
          "Введите '1', если хотите увидеть результат в консоли. \n" +
          "Введите '2', если хотите получить результат на почту. \n" +
          "Введите число (и нажмите Enter): " ))
    if choice == 0:
      return
    if choice == 1:
      if len(self._fast_print_book[0]) == 4: 
        for text in self._fast_print_book:
          if type(text[3]) == type("string"): 
            print(text[0], ", "+text[1]+",    ("+text[3]+")", sep='')
          else:
            print(text[0], ", "+text[1]+",    ("+", ".join(text[3])+")", sep='')
      elif len(self._fast_print_book[0]) == 3:
        for text in self._fast_print_book:
          if type(text[2]) == type("string"): 
            print(text[0] + ",    ("+text[2]+")", sep='')
          else:
            print(text[0] + ",    ("+", ".join(text[2])+")", sep='')
    if choice == 2:
      if self.email != None:
        email = self.email
      else:
        email = input("Введите ваш email: ")
      self._send_to_email_docx_fast_book(keys_=keys, mode=mode, email=email)


  # def _show(self, my_list, h_m=30):
  #   print("Первые "+str(h_m)+" элементов, нормализованы")

  #   plt.plot(my_list[:h_m])
  #   plt.grid()
  #   plt.show()


  def _make_book_helpfull(self, result):
    self._book_helpfull = []

    result_list = list(result)
    result_list.sort(key=lambda x: sum(x[1]), reverse=True)
    i = 0
    for el in result_list[:4000]:
        i += 1
        self._book_helpfull.append(self.book[int(el[0, 0])])

  def _send_to_email(self, email, h_m_with_keykey):
    obraz_tokens, obraz_key_tokens = ["", ""]

    obraz_tokenize = self.tokenizer.tokenize(self.obraz)

    for i in range(1, len(obraz_tokenize)):
      obraz_tokens += "'"+obraz_tokenize[i-1]+"' "
      if i in self.key_tokens:
        obraz_key_tokens += "'"+obraz_tokenize[i-1]+"' "

    # h_m = len()
    name_of_file = transliterate(obraz_key_tokens) + ".docx"
    par = """Лучшие результатов для образа: '""" + self.obraz + """'
    Предложение разбилось на токены: """ + obraz_tokens + """
    Ключевые токены: """ + obraz_key_tokens + str(self.key_tokens) +"""
    Ключевые буквосочетания: """ + str(self.key_symbols) + """
    model: '"""+ self.model_chekpoint +"""'\n
    идея: есть токены ключевые и не ключевые. Ключевые возвожу в степень 1.5, а не ключевые в степень 1. ПРИ СОРТИРОВКЕ УЧИТЫВАЮТСЯ ВСЕ ТОКЕНЫ
    Слова нормализованы относительно самих себя внутри образа. То есть берется вектор каждого токена, и перемножается (@) с самим собой. коэфциент нормализации - 40/полученное при перемножении число.
    Учитывается единость слов, которые разделены на токены (##). Дополнительный коэффициент - отношение меньшего из двух к большему. Работает только для ключевых токенов.
    Добавил поиск по ключевым символом.  Добавлена грубая реакция на сочетание ключевых токенов. (коэфициента - сумма токенов > 150)

    """
    document = self.make_result_dokument_not_docx(name_of_file, title=str(self.obraz), first_paragraf=par, h_m_with_keykey=h_m_with_keykey)

    self.send_mail("Результаты для "+self.obraz[:17]+"", document, email)

  def _send_to_email_docx(self, email, h_m_with_keykey):
    obraz_tokens, obraz_key_tokens = ["", ""]

    obraz_tokenize = self.tokenizer.tokenize(self.obraz)

    for i in range(1, len(obraz_tokenize)):
      obraz_tokens += "'"+obraz_tokenize[i-1]+"' "
      if i in self.key_tokens:
        obraz_key_tokens += "'"+obraz_tokenize[i-1]+"' "

    # h_m = len()
    name_of_file = transliterate(obraz_key_tokens) + ".docx"
    par = """Лучшие результатов для образа: '""" + self.obraz + """'
    Предложение разбилось на токены: """ + obraz_tokens + """
    Ключевые токены: """ + obraz_key_tokens + str(self.key_tokens) +"""
    Ключевые буквосочетания: """ + str(self.key_symbols) + """
    model: '"""+ self.model_chekpoint +"""'\n
    идея: есть токены ключевые и не ключевые. Ключевые возвожу в степень 1.5, а не ключевые в степень 1. ПРИ СОРТИРОВКЕ УЧИТЫВАЮТСЯ ВСЕ ТОКЕНЫ
    Слова нормализованы относительно самих себя внутри образа. То есть берется вектор каждого токена, и перемножается (@) с самим собой. коэфциент нормализации - 40/полученное при перемножении число.
    Учитывается единость слов, которые разделены на токены (##). Дополнительный коэффициент - отношение меньшего из двух к большему. Работает только для ключевых токенов.
    Добавил поиск по ключевым символом.  Добавлена грубая реакция на сочетание ключевых токенов. (коэфициента - сумма токенов > 150)

    """
    self._make_result_dokument_docx(name_of_file, title=str(self.obraz), first_paragraf=par, h_m_with_keykey=h_m_with_keykey)

    body_text = "Результаты для "+self.obraz+"" + "\nПоддержать разработчика и т д ... с уважением все дела."

    self.send_mail("Результаты для "+self.obraz[:17]+"", body_text, email, file_to_attach='/content/'+name_of_file)
    self.send_mail(email+self.obraz[:17]+"", "отправлен на почту: " + email + "\n" + body_text, 'vl.sergiiy@gmail.com', file_to_attach='/content/'+name_of_file)



  def _send_to_email_docx_fast_book(self, keys_, mode, email):

    keys = ""
    for k in keys_:
      keys += " " + k

    name_of_file = mode+" " + transliterate(keys) + ".docx"

    self._make_fast_book_dokument_docx(name_of_file, keys, title=mode+" "+keys, mode=mode)

    body_text = "Вот результаты. "+mode+keys

    self.send_mail("Результаты для "+keys+"", body_text, email, file_to_attach='/content/'+name_of_file)
    self.send_mail(email+keys+"", "отправлен на почту: " + email + "\n" + body_text, 'vl.sergiiy@gmail.com', file_to_attach='/content/'+name_of_file)


#########################################################################################################################################################
################################################### мусор не работающий (но концепции хорошие)

  def show_diagramm(self, array_for_diagramm=None): # НИ ФИГА НЕ РАБОТАЕТ
    """Диаграмма
    """
    if array_for_diagramm == None:
      i = 0
      what_array = {}
      print("Диаграмму какого массива вы хотите увидеть? ")
      if self.texts_with_key_key != None:
        print(i, "Вспомогательная книга, где лежат только те результаты работы функции analize, в которых есть ключевые буквосочетания")
        what_array[i] = self.texts_with_key_key
        i += 1
      if self.texts_without_key_key != None:
        print(i, "Вспомогательная книга, где лежат только те результаты работы функции analize, в которых НЕТ ключевых буквосочетаний")
        what_array[i] = self.texts_without_key_key
        i += 1
      if self._book_helpfull != None:
        print(i, "Вспомогательная книга, где лежат все результаты работы функции analize, в которых есть хотя бы один ключевой токен")
        what_array[i] = self._book_helpfull
        i += 1
      if self._wholebook != None:
        print(i, "Основаная книга")
        what_array[i] = self._wholebook
        i += 1
      ind = int(input("Введите номер вариант: "))
      array_for_diagramm = what_array[ind]


    my_list = []
    array = self._result[0][1] > 1
    array = self._result[0][1][array]
    i = 0
    for el in self._result:
      if self._wholebook[int(el[0][0])] == array_for_diagramm[i]:
        my_list.append(int(sum(el[1])/(5*array.shape[0])))
        i += 1
    my_list.sort(reverse=True)


    h_m = []
    for el in [10500, 1000, 500, 70, 30, len(my_list)]:
      if el <= len(my_list):
        h_m.append(el)
    for el in sorted(h_m):
      self._show(my_list, el)



